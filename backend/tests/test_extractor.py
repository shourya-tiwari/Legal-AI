import io

import docx

from app.services.extractor import extract_text_and_blocks


def test_txt_extraction_returns_full_text_and_blocks():
    content = b"Paragraph one.\n\nParagraph two."
    result = extract_text_and_blocks(content, "contract.txt", "text/plain")

    assert "Paragraph one." in result["full_text"]
    assert "Paragraph two." in result["full_text"]
    assert len(result["blocks"]) == 2
    assert result["blocks"][0]["id"] == 1
    assert result["blocks"][0]["type"] == "paragraph"
    assert result["blocks"][0]["page"] == 1


def test_txt_extraction_empty_input_returns_no_blocks():
    result = extract_text_and_blocks(b"", "empty.txt", "text/plain")

    assert result["full_text"] == ""
    assert result["blocks"] == []


def test_docx_extraction_returns_paragraphs():
    document = docx.Document()
    document.add_paragraph("This Agreement is made on January 1, 2025.")
    document.add_paragraph("The Tenant shall pay rent monthly.")
    buf = io.BytesIO()
    document.save(buf)

    result = extract_text_and_blocks(
        buf.getvalue(), "lease.docx",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    assert "This Agreement is made on January 1, 2025." in result["full_text"]
    assert len(result["blocks"]) == 2
    assert all(b["page"] == 1 for b in result["blocks"])


def test_unknown_extension_falls_back_to_plain_text_decode():
    content = b"Just some plain content with no known extension."
    result = extract_text_and_blocks(content, "notes.xyz", None)

    assert "Just some plain content" in result["full_text"]
    assert len(result["blocks"]) == 1
